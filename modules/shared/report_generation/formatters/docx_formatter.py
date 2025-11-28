"""
DOCX FORMATTER - DOCX Format Export

Exports reports to DOCX (.docx) format with:
- Editable format
- Professional styling
- Court-ready output
- Metadata

Features:
- Comprehensive docstrings
- Error handling with custom exceptions
- Structured logging
- Graceful degradation (fallback to text if DOCX library unavailable)
"""

import logging
import json
from typing import Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)

class FormatterException(Exception):
    """Base exception for formatter errors"""
    pass

class FormattingError(FormatterException):
    """Raised when formatting fails"""
    pass

class StructuredLogger:
    """Structured logging with JSON context"""
    
    @staticmethod
    def log_with_context(level: str, message: str, **context) -> None:
        """Log with context information"""
        try:
            log_entry = {
                'timestamp': datetime.now().isoformat(),
                'level': level,
                'message': message,
                'context': context
            }
            log_level = getattr(logging, level.upper(), logging.INFO)
            logger.log(log_level, json.dumps(log_entry))
        except Exception as e:
            logger.error(f"Error in structured logging: {str(e)}")

class DOCXFormatter:
    """
    Format reports to DOCX (.docx) format.
    
    Provides professional DOCX formatting with editable content,
    professional styling, and court-ready output. Includes graceful
    degradation if DOCX libraries are not available.
    """
    
    def __init__(self):
        """Initialize DOCX formatter"""
        logger.debug("DOCXFormatter initialized")
        self.docx_available = self._check_docx_library()
    
    @staticmethod
    def _check_docx_library() -> bool:
        """
        Check if DOCX library is available.
        
        Returns:
            bool: True if DOCX library available, False otherwise
        """
        try:
            from docx import Document
            return True
        except ImportError:
            logger.warning("python-docx not available, DOCX export will use fallback")
            return False
    
    @staticmethod
    def format(report_content: str, case_id: str = "") -> str:
        """
        Format report content to DOCX.
        
        Converts report content to DOCX format with professional styling,
        editable content, and metadata. Falls back to text format if DOCX
        libraries are not available.
        
        Args:
            report_content (str): Report content to format
            case_id (str): Case ID for logging (optional)
            
        Returns:
            str: DOCX content or fallback text content
            
        Raises:
            FormattingError: If formatting fails
            
        Example:
            >>> formatter = DOCXFormatter()
            >>> docx_content = formatter.format(report_content, "CASE-001")
            >>> len(docx_content) > 0
            True
        """
        try:
            structured_logger = StructuredLogger()
            structured_logger.log_with_context(
                "DEBUG",
                "Formatting report to DOCX",
                case_id=case_id,
                content_length=len(report_content)
            )
            
            # Try to use python-docx
            try:
                from docx import Document
                from io import BytesIO
                
                # Create DOCX document
                doc = Document()
                
                # Add title
                doc.add_heading(f'Forensic Report - {case_id}', 0)
                
                # Add content
                for line in report_content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                
                # Save to bytes
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                
                structured_logger.log_with_context(
                    "DEBUG",
                    "DOCX formatting completed",
                    case_id=case_id,
                    docx_size=docx_buffer.tell()
                )
                
                return docx_buffer.getvalue().decode('latin-1', errors='ignore')
            
            except ImportError:
                logger.warning("python-docx not available, returning text format")
                structured_logger.log_with_context(
                    "WARNING",
                    "DOCX library not available, using text fallback",
                    case_id=case_id
                )
                return report_content
        
        except Exception as e:
            error_msg = f"Error formatting to DOCX: {str(e)}"
            logger.error(error_msg)
            raise FormattingError(error_msg) from e
